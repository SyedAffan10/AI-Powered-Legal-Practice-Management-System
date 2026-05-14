from googleapiclient.http import MediaIoBaseDownload
from googleapiclient.discovery import build
from google.oauth2.service_account import Credentials
from googleapiclient.http import MediaIoBaseDownload
from openai import OpenAI
from PyPDF2 import PdfReader
from docx import Document
from dotenv import load_dotenv
import time
import io
import os
import PyPDF2

load_dotenv()

api_key = os.getenv('OPENAI_API_KEY')
client = OpenAI(api_key=api_key)

# Create Drive service
SERVICE_ACCOUNT_FILE = 'service_account_credentials.json'
# SERVICE_ACCOUNT_FILE = 'new_credentials.json'
drive_service = Credentials.from_service_account_file(
    SERVICE_ACCOUNT_FILE, scopes=['https://www.googleapis.com/auth/drive']
)
service = build('drive', 'v3', credentials=drive_service)

# Function to fetch DOCX content from Google Drive
def fetch_doc_content(file_id):
    request = service.files().get_media(fileId=file_id)
    file_data = io.BytesIO()
    downloader = MediaIoBaseDownload(file_data, request)

    done = False
    while not done:
        status, done = downloader.next_chunk()
        print(f"Download progress: {int(status.progress() * 100)}%")

    file_data.seek(0)
    doc_content = []
    document = Document(file_data)
    for para in document.paragraphs:
        doc_content.append(para.text)

    return "\n".join(doc_content)


def fetch_pdf_content(file_id):
    request = service.files().get_media(fileId=file_id)
    file_data = io.BytesIO()
    downloader = MediaIoBaseDownload(file_data, request)

    done = False
    while not done:
        status, done = downloader.next_chunk()
        print(f"Download progress: {int(status.progress() * 100)}%")

    file_data.seek(0)
    pdf_content = []
    reader = PdfReader(file_data)
    for page in reader.pages:
        pdf_content.append(page.extract_text())

    return "\n".join(pdf_content)


MAX_TOKENS = 15000  # Adjust based on OpenAI token limits
CHUNK_SIZE = 10000  # Split large content into manageable chunks

def process_plan(schools_data, template_file, prompt=None, add_files=None):
    all_messages = {}

    schools = schools_data.get("schools", {})

    for school_name, school_data in schools.items():
        print(f"Processing school: {school_name}")

        files = school_data.get("files", {})
        docx_files = files.get("docx_files", [])
        pdf_files = files.get("pdf_files", [])

        # Fetch document content
        doc_content = fetch_doc_content(docx_files[0]["file_id"]) if docx_files else ""
        template_content = fetch_doc_content(template_file["file_id"])

        coach_pdf = None
        other_pdfs = []

        for pdf in pdf_files:
            file_name = pdf.get("file_name", "").lower()

            if "coach" in file_name:
                coach_pdf = fetch_pdf_content(pdf["file_id"])
            else:
                other_pdfs.append(fetch_pdf_content(pdf["file_id"]))

        # Handling different PDF combinations
        pdf1_content, pdf2_content, pdf3_content = "", "", ""
        if coach_pdf and len(other_pdfs) == 1:
            pdf1_content = coach_pdf
            pdf2_content = other_pdfs[0]
        elif coach_pdf and len(other_pdfs) == 0:
            pdf1_content = coach_pdf
        elif not coach_pdf and len(other_pdfs) == 1:
            pdf2_content = other_pdfs[0]
        elif not coach_pdf and len(other_pdfs) > 1:
            pdf2_content = other_pdfs[0]
            pdf3_content = other_pdfs[1]
        elif coach_pdf and len(other_pdfs) > 1:
            pdf1_content = coach_pdf
            pdf2_content = other_pdfs[0]
            pdf3_content = other_pdfs[1]

        # Clean the contents with GPT
        cleaned_pdf1_content = safe_gpt_request(pdf1_content)
        print("Perfect - clean_and_organize_content_with_gpt_1")

        cleaned_pdf2_content = safe_gpt_request(pdf2_content)
        print("Perfect - clean_and_organize_content_with_gpt_2")

        cleaned_pdf3_content = safe_gpt_request(pdf3_content)
        print("Perfect - clean_and_organize_content_with_gpt_3")

        cleaned_template_content = clean_and_organize_template_with_gpt(template_content)
        print("Perfect - clean_and_organize_template_with_gpt")

        # Generate final message
        message = generate_message_gpt(
            doc_content, cleaned_pdf1_content, cleaned_pdf2_content, cleaned_pdf3_content,
            cleaned_template_content, prompt, add_files
        )
        print("Perfect - generate_message_gpt")

        all_messages[school_name] = message

        time.sleep(5)

    return all_messages



def safe_gpt_request(content):
    """Handles large content, rate limits, and retries GPT calls."""
    if not content:
        return ""

    # Trim content if too large
    content = trim_content(content, MAX_TOKENS)

    # Split into chunks if still large
    chunks = split_text_into_chunks(content, CHUNK_SIZE)

    cleaned_chunks = []
    for chunk in chunks:
        try:
            cleaned_chunks.append(clean_and_organize_content_with_gpt(chunk))
        except Exception as e:
            print(f"Rate limit error: {e}")
            time.sleep(5)  # Wait 5 seconds before retrying
            try:
                cleaned_chunks.append(clean_and_organize_content_with_gpt(chunk))
            except Exception as e:
                print(f"Failed after retry: {e}")
                cleaned_chunks.append("")  # Add empty string on failure

    return "\n".join(cleaned_chunks)  # Combine processed chunks


def trim_content(content, max_tokens=MAX_TOKENS):
    """Trims content to a safe length before sending to GPT."""
    words = content.split()
    if len(words) > max_tokens:
        print(f"Trimming content from {len(words)} tokens to {max_tokens} tokens.")
        return " ".join(words[:max_tokens])
    return content


def split_text_into_chunks(text, max_length=CHUNK_SIZE):
    """Splits large text into smaller chunks."""
    words = text.split()
    return [" ".join(words[i:i + max_length]) for i in range(0, len(words), max_length)]


def clean_and_organize_content_with_gpt(content):
    """Cleans content using GPT, ensuring markdown output and handling errors."""
    if not content:
        return ""
    
    prompt = (
        "You are a content-cleaning assistant. Your task is to clean and organize the following content. "
        "Remove irrelevant information like standalone numbers, percentages, spaces and symbols that do not add meaning. "
        "Structure the text into meaningful and concise sentences. "
        "**Return the final output in Markdown format.**"
    )
    gpt_prompt = f"{prompt}\n\n**Content**:\n{content}\n\nPlease return the cleaned and organized content in Markdown format."
    response = client.chat.completions.create(
        model="gpt-4o",
        messages=[
            {"role": "system", "content": "You are a content-cleaning assistant."},
            {"role": "user", "content": gpt_prompt}
        ],
        temperature=0.0
    )
    return response.choices[0].message.content


def clean_and_organize_template_with_gpt(content):
    if not content:
        print("Error: content is empty!")
        return ""

    prompt = (
        "You are a content-organizing assistant. Your task is to format the given content exactly as provided, "
        "but with the specified headings in **bold**. Do not alter the structure, spacing, or paragraph formatting of the content.\n\n"
        "**Headings to be bold:**\n"
        "- **<College> <sport>**"
        "- **<Month> <Year>**"
        "- **<Month>/<Month> <Year>**"
        "- **<Month>/<Month>/<Month> <Year>**"
        "- **<Month>/<Month>/<Month>/<Month> <Year>**"
        "- **TRS Messages**"
        "- **<Month>: <Topic>**"
        "- **Talking Points**"
        "- **Social Media Topic Ideas**"
        "- **Text Message Talking Points**"
        "- **Week <number>**"
        "- **WEEK <number>**"
        "- **Email <number>**"
        "- **Letter <number>**"
        "- **Parent Letter**"
        "- **Coach Letter**"
        "- **Coach Letter or Email**"
        "- **Visit Letter to Parents**"
        "- **Portal Elite Messaging**"
        "- **Use anytime in <Month> or <Month>**\n\n"
        "**Additional Formatting Rule:**\n"
        "- Any **month** or **combination of months** followed by a **year** should be in **bold** (e.g., Jan 2025, Feb/Mar 2025, Feb./Mar. 2025).\n"
        "- Any **month followed by a colon and a topic** should be in **bold** (e.g., March: Athletic Atmosphere, April: Coaching).\n\n"
        "If you identify any other headings in the content, format them in **bold** as well.\n\n"
        "Additionally, apply only bulleted points where appropriate to improve readability. Do not use numbered lists.\n\n"
        "**Content:**\n{content}\n\n"
        "**Final Output:**\n"
        "Return the content exactly as provided, preserving the original spacing and structure. Ensure the specified headings and any other detected headings are in **bold**. "
        "Apply only bulleted points where necessary to enhance clarity, without introducing numbered lists."
    )

    gpt_prompt = prompt.format(content=content)
    response = client.chat.completions.create(
        model="gpt-4o",
        messages=[
            {"role": "system", "content": "You are a content-organizing assistant."},
            {"role": "user", "content": gpt_prompt}
        ],
        temperature=0.0
    )
    cleaned_content = response.choices[0].message.content
    return cleaned_content


def generate_message_gpt(doc_content, pdf1_content, pdf2_content, pdf3_content, template_text, prompt=None, add_files=None):
    """
    Generate mail content using DOCX and PDF files dynamically, including uploaded files in the prompt.
    """
    # Initialize placeholders for additional content
    additional_docx_content = ""
    additional_pdf_content = []

    # Process uploaded files directly from 'add_files'
    if add_files:
        for file in add_files:
            file_name = file.filename
            file_content = file.file.read()
            
            # Check file type and process accordingly
            if file_name.endswith(".pdf"):
                # Extract text from PDF
                try:
                    pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
                    pdf_text = ""
                    for page in pdf_reader.pages:
                        pdf_text += page.extract_text() or ""
                    additional_pdf_content.append(pdf_text.strip())
                except Exception as e:
                    additional_pdf_content.append(f"Error processing {file_name}: {str(e)}")
            
            elif file_name.endswith(".docx"):
                # Extract text from DOCX
                try:
                    document = Document(io.BytesIO(file_content))
                    docx_text = "\n".join([paragraph.text for paragraph in document.paragraphs])
                    additional_docx_content += docx_text + "\n"
                except Exception as e:
                    additional_docx_content += f"Error processing {file_name}: {str(e)}\n"

    # Combine PDF and DOCX content into strings
    combined_additional_pdf_content = "\n".join(additional_pdf_content).strip()
    combined_additional_docx_content = additional_docx_content.strip()

    # Ensure that if content is empty, we replace it with an empty string ""
    if not combined_additional_pdf_content:
        combined_additional_pdf_content = '""'
    if not combined_additional_docx_content:
        combined_additional_docx_content = '""'

    base_prompt = f"""
The following content has been cleaned and structured. It includes details extracted from multiple documents, primarily containing reviews from coaches and athletes. Use this refined content to generate a well-structured response:

- **DOCX Content**: {doc_content}  
- **Coach Survey**: {pdf1_content} (Feedback from coaches)  
- **Athlete Survey 1**: {pdf2_content} (Feedback from athletes)  
- **Athlete Survey 2**: {pdf3_content} (Feedback from athletes)  
- **Additional PDF Content**: {combined_additional_pdf_content}  
- **Additional DOCX Content**: {combined_additional_docx_content}  

### Instructions:
1. Replace all placeholders in the template below with concise and relevant information derived from the content above.
2. **Ensure the following:**  
    - If a placeholder cannot be filled due to missing or unclear information, leave it unchanged (e.g., '<placeholder>').  
    - Use **bold headings** where applicable.  
    - Structure the response with bulleted points to improve readability.  
    - Prioritize **Coach Survey feedback** over Athlete Surveys in case of conflicting information.  
    - Do **not** alter the original meaning of the content — only organize and structure it effectively.

### Template:
{template_text}

**IMPORTANT**: Only return the fully completed template. Do not omit any sections or placeholders. Do not include additional notes, comments, extra spaces, or extraneous information. Ensure that **bold headings** and structured lists (bulleted) are used where appropriate.
"""

    final_prompt = f"{base_prompt}\n\nPlease understand this input carefully and rewrite the content accordingly:\n{prompt}" if prompt else base_prompt

    response = client.chat.completions.create(
        model="gpt-4o",
        messages=[ 
            {"role": "system", "content": "You are a content generation assistant specialized in contextual replacements."},
            {"role": "user", "content": final_prompt}
        ],
        temperature=0.0
    )

    filled_template = response.choices[0].message.content
    return filled_template