import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
from googleapiclient.http import MediaIoBaseUpload

def extract_text_from_dict(data):
    all_text = []
    if isinstance(data, dict):
        for value in data.values():
            all_text.append(extract_text_from_dict(value))
    elif isinstance(data, str):
        all_text.append(data)
    return "\n".join(all_text)

def extract_headings(paragraphs):
    headings = []
    for para in paragraphs:
        # Adjust the regex to handle both patterns: "- **" and "**"
        match = re.match(r'^\s*- \*\*(.*)\*\*\s*$', para.strip())  # Matches "- **heading**"
        if match:
            headings.append(match.group(1))  # Extract the heading text without the "- **" and "**"
        else:
            match = re.match(r'^\*\*(.*)\*\*\s*$', para.strip())  # Matches "**heading**"
            if match:
                headings.append(match.group(1))  # Extract the heading text without the "**"
    return headings


def add_logo_to_page(doc, logo_path):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    run.add_picture(logo_path, width=Inches(2))
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT


def process_and_save_colleges_files(colleges_data, save_folder_id, service, logo_path="./logo/logos_proj.jpeg"):
    """
    Save generated files directly to Google Drive, structured by group and schools.
    """
    def create_folder(folder_name, parent_id):
        """Create a folder in Google Drive."""
        try:
            folder_metadata = {
                "name": folder_name,
                "mimeType": "application/vnd.google-apps.folder",
                "parents": [parent_id]
            }
            folder = service.files().create(body=folder_metadata, fields="id").execute()
            return folder.get("id")
        except Exception as e:
            print(f"Error creating folder {folder_name}: {e}")
            raise
    
    uploaded_files = []

    for group_name, schools in colleges_data.items():
        # Create a group folder
        group_folder_id = create_folder(group_name, save_folder_id)

        for school_name, school_data in schools.items():
            # Create a school subfolder
            school_folder_id = create_folder(school_name, group_folder_id)

            # Initialize the document for the school
            doc = Document()

            # Add logo
            add_logo_to_page(doc, logo_path)

            # Extract all textual content
            result1 = extract_text_from_dict(school_data)
            lines = result1.splitlines()

            # Extract and process headings
            all_headings = extract_headings(lines)
            print("headings: ", all_headings)

            if len(all_headings) >= 3:
                first_heading = all_headings[0]
                second_heading = all_headings[1]
                third_heading = all_headings[2]
                headings = all_headings[3:]
            else:
                print("Not enough headings found in the document.")
                first_heading = second_heading = third_heading = ""
                headings = []

            # Add main headings
            if first_heading:
                heading = doc.add_heading(first_heading, level=1)
                heading_font = heading.runs[0].font
                heading_font.size = Pt(24)
                heading_font.color.rgb = RGBColor(11, 83, 148)
                heading.paragraph_format.line_spacing = 1
                heading.paragraph_format.space_before = Pt(0)
                heading.paragraph_format.space_after = Pt(6)
                heading_font.name = 'Arial'
                heading_font.bold = True

            if second_heading:
                heading = doc.add_heading(second_heading, level=1)
                heading_font = heading.runs[0].font
                heading_font.size = Pt(24)
                heading_font.color.rgb = RGBColor(230, 145, 56)
                heading.paragraph_format.line_spacing = 1
                heading.paragraph_format.space_before = Pt(0)
                heading.paragraph_format.space_after = Pt(6)
                heading_font.name = 'Arial'
                heading_font.bold = True

            if third_heading:
                heading = doc.add_heading(third_heading, level=1)
                heading_font = heading.runs[0].font
                heading_font.size = Pt(24)
                heading_font.color.rgb = RGBColor(230, 145, 56)
                heading.paragraph_format.line_spacing = 1
                heading.paragraph_format.space_before = Pt(0)
                heading.paragraph_format.space_after = Pt(6)
                heading_font.name = 'Arial'
                heading_font.bold = True

            collecting_content = True
            current_heading = None

            for line in lines:
                line_text = line.strip()

                # Check if the line matches either of the heading patterns
                if re.match(r'^\s*(- \*\*.*\*\*\s*|\*\*.*\*\*\s*)$', line_text):
                    heading_text = line_text.strip('- **').strip('**')  # Remove "- **" and "**"

                    if heading_text in [first_heading, second_heading, third_heading]:
                        collecting_content = True
                        current_heading = heading_text
                        continue

                    if heading_text.lower().startswith("week"):
                        heading = doc.add_heading(heading_text, level=1)
                        heading_font = heading.runs[0].font
                        heading_font.size = Pt(16)
                        heading_font.color.rgb = RGBColor(11, 83, 148)
                        heading_font.name = 'Arial'
                        heading.paragraph_format.line_spacing = 1
                        heading.paragraph_format.space_before = Pt(0)
                        heading.paragraph_format.space_after = Pt(6)
                        current_heading = heading_text
                        continue

                    keywords = ["email", "letter", "parent", "coach", "send", "suggested", "if", "visit"]
                    if any(heading_text.lower().startswith(keyword.lower()) for keyword in keywords):
                        heading = doc.add_heading(heading_text, level=2)
                        heading_font = heading.runs[0].font
                        heading_font.size = Pt(16)
                        heading_font.color.rgb = RGBColor(230, 145, 56)
                        heading_font.name = 'Arial'
                        heading.paragraph_format.line_spacing = 1
                        heading.paragraph_format.space_before = Pt(0)
                        heading.paragraph_format.space_after = Pt(6)
                        current_heading = heading_text
                        continue

                    keywords = ["visit"]
                    if current_heading and any(current_heading.lower().startswith(keyword.lower()) for keyword in keywords):
                        heading = doc.add_heading(heading_text, level=1)
                        heading_font = heading.runs[0].font
                        heading_font.size = Pt(16)
                        heading_font.color.rgb = RGBColor(230, 145, 56)
                        heading_font.name = 'Arial'
                        heading.paragraph_format.line_spacing = 1
                        heading.paragraph_format.space_before = Pt(0)
                        heading.paragraph_format.space_after = Pt(6)
                        current_heading = heading_text
                        collecting_content = True
                        continue

                    valid_start_keywords = ('talking', 'social', 'text')
                    if heading_text.lower().startswith(valid_start_keywords):
                        heading = doc.add_heading(heading_text, level=1)
                        heading_font = heading.runs[0].font
                        heading_font.size = Pt(16)
                        heading_font.color.rgb = RGBColor(230, 145, 56)
                        heading_font.name = 'Arial'
                        heading.paragraph_format.line_spacing = 1
                        heading.paragraph_format.space_before = Pt(0)
                        heading.paragraph_format.space_after = Pt(6)
                        current_heading = heading_text
                    else:
                        doc.add_page_break()
                        add_logo_to_page(doc, logo_path)
                        heading = doc.add_heading(heading_text, level=1)
                        heading_font = heading.runs[0].font
                        heading_font.size = Pt(16)
                        heading_font.color.rgb = RGBColor(11, 83, 148)
                        heading_font.name = 'Arial'
                        heading.paragraph_format.line_spacing = 1
                        heading.paragraph_format.space_before = Pt(0)
                        heading.paragraph_format.space_after = Pt(6)
                        current_heading = heading_text

                else:
                    if collecting_content:
                        if line.startswith("-"):
                            paragraph = doc.add_paragraph(line[1:].strip(), style='List Bullet')
                            paragraph.paragraph_format.line_spacing = 1
                            paragraph.paragraph_format.space_before = Pt(0)
                            if paragraph.runs:
                                paragraph_font = paragraph.runs[0].font
                                paragraph_font.name = 'Arial'
                                paragraph_font.size = Pt(12)

                        elif line.lstrip().startswith("-"):  # Sub-bullets, considering spaces before the dash
                            paragraph = doc.add_paragraph(line.lstrip()[1:].strip(), style='List Bullet 2')
                            paragraph.paragraph_format.line_spacing = 1
                            paragraph.paragraph_format.space_before = Pt(0)
                            if paragraph.runs:
                                paragraph_font = paragraph.runs[0].font
                                paragraph_font.name = 'Arial'
                                paragraph_font.size = Pt(12)

                        else:
                            paragraph = doc.add_paragraph(line_text)
                            paragraph.paragraph_format.line_spacing = 1
                            paragraph.paragraph_format.space_before = Pt(0)
                            if paragraph.runs:
                                paragraph_font = paragraph.runs[0].font
                                paragraph_font.name = 'Arial'
                                paragraph_font.size = Pt(12)

            # Save the document to a BytesIO buffer
            doc_stream = BytesIO()
            doc.save(doc_stream)
            doc_stream.seek(0)

            try:
                file_metadata = {
                    "name": f"{school_name}.docx",
                    "parents": [school_folder_id]
                }
                media = MediaIoBaseUpload(doc_stream, mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
                uploaded_file = service.files().create(body=file_metadata, media_body=media, fields="id, name, parents").execute()
                uploaded_files.append(uploaded_file)
                print(f"Uploaded {school_name}.docx to Google Drive under {group_name}.")
            except Exception as e:
                print(f"Error uploading file {school_name}.docx: {e}")

    return uploaded_files
