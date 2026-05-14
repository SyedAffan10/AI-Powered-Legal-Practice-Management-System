import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
from googleapiclient.http import MediaIoBaseUpload

def extract_text_from_dict(data):
    all_text = []
    if isinstance(data, dict):
        for value in data.values():
            all_text.append(extract_text_from_dict(value))
    elif isinstance(data, str):
        all_text.append(data)
    return "\n".join(all_text)

def extract_headings(paragraphs):
    headings = []
    for para in paragraphs:
        match = re.match(r'^\s*- \*\*(.*)\*\*\s*$', para.strip())
        if match:
            headings.append(match.group(1))
        else:
            match = re.match(r'^\*\*(.*)\*\*\s*$', para.strip())
            if match:
                headings.append(match.group(1))
    return headings

def add_logo_to_page(doc, logo_path):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    run.add_picture(logo_path, width=Inches(2))
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

def add_heading(doc, text, level, font_size, color_rgb):
    heading = doc.add_heading(text, level=level)
    run = heading.runs[0]
    run.font.size = Pt(font_size)
    run.font.color.rgb = RGBColor(*color_rgb)
    run.font.name = 'Arial'
    heading.paragraph_format.line_spacing = 1
    heading.paragraph_format.space_before = Pt(0)
    heading.paragraph_format.space_after = Pt(6)
    return heading

def add_paragraph(doc, text, style=None):
    paragraph = doc.add_paragraph(text, style=style) if style else doc.add_paragraph(text)
    paragraph.paragraph_format.line_spacing = 1
    paragraph.paragraph_format.space_before = Pt(0)
    if paragraph.runs:
        paragraph.runs[0].font.name = 'Arial'
        paragraph.runs[0].font.size = Pt(12)
    return paragraph

def process_and_save_portal_elite_colleges_files(colleges_data, save_folder_id, service, logo_path="./logo/logos_proj.jpeg"):
    """
    Save generated files directly to Google Drive, structured by group and schools.
    """
    def create_folder(folder_name, parent_id):
        """Create a folder in Google Drive."""
        try:
            folder_metadata = {
                "name": folder_name,
                "mimeType": "application/vnd.google-apps.folder",
                "parents": [parent_id]
            }
            folder = service.files().create(body=folder_metadata, fields="id").execute()
            return folder.get("id")
        except Exception as e:
            print(f"Error creating folder {folder_name}: {e}")
            raise
    
    uploaded_files = []

    for group_name, schools in colleges_data.items():
        group_folder_id = create_folder(group_name, save_folder_id)

        for school_name, school_data in schools.items():
            school_folder_id = create_folder(school_name, group_folder_id)

            # Initialize the document for the school
            doc = Document()

            full_text = extract_text_from_dict(school_data)

            first_week_index = full_text.find("**Week 1:")
            second_week_index = full_text.find("**Week 1:", first_week_index + 1)

            intro_text = full_text[:first_week_index].strip()
            if second_week_index != -1:
                weekly_text = full_text[first_week_index:second_week_index].strip()
                detailed_text = full_text[second_week_index:].strip()
            else:
                weekly_text = full_text[first_week_index:].strip()
                detailed_text = ""

            intro_lines = intro_text.splitlines()
            intro_headings = extract_headings(intro_lines)
            if len(intro_headings) >= 2:
                first_heading = intro_headings[0]
                second_heading = intro_headings[1]
            else:
                first_heading = second_heading = ""

            doc = Document()
            add_logo_to_page(doc, logo_path)

            if first_heading:
                add_heading(doc, first_heading, level=1, font_size=24, color_rgb=(11, 83, 148))
            if second_heading:
                add_heading(doc, second_heading, level=1, font_size=24, color_rgb=(230, 145, 56))

            for line in intro_lines:
                line_text = line.strip()
                if line_text.startswith("**") and (first_heading in line_text or second_heading in line_text):
                    continue
                if not line_text:
                    continue
                add_paragraph(doc, line_text)

            doc.add_page_break()
            add_logo_to_page(doc, logo_path)

            weekly_lines = weekly_text.splitlines()
            for line in weekly_lines:
                line_text = line.strip()
                if not line_text:
                    continue
                if re.match(r'^\s*(- \*\*.*\*\*\s*|\*\*.*\*\*\s*)$', line_text):
                    heading_text = line_text.strip('- **').strip('**')
                    if heading_text.lower().startswith("week"):
                        add_heading(doc, heading_text, level=1, font_size=16, color_rgb=(11, 83, 148))
                    else:
                        add_heading(doc, heading_text, level=2, font_size=16, color_rgb=(230, 145, 56))
                else:
                    add_paragraph(doc, line_text)

            if detailed_text:
                detailed_lines = detailed_text.splitlines()
                for line in detailed_lines:
                    line_text = line.strip()
                    if not line_text:
                        continue
                    if re.match(r'^\s*(- \*\*.*\*\*\s*|\*\*.*\*\*\s*)$', line_text):
                        heading_text = line_text.strip('- **').strip('**')
                        is_week_heading = heading_text.lower().startswith("week")
                        is_subheading = any(heading_text.lower().startswith(k) for k in ["talking", "social", "text"])
                        is_email_heading = any(heading_text.lower().startswith(k) for k in ["email", "letter", "parent", "coach", "send", "suggested", "if", "visit"])
                        if is_week_heading:
                            doc.add_page_break()
                            add_logo_to_page(doc, logo_path)
                            add_heading(doc, heading_text, level=1, font_size=16, color_rgb=(11, 83, 148))
                        elif is_subheading or is_email_heading:
                            level = 2 if is_email_heading else 1
                            add_heading(doc, heading_text, level=level, font_size=16, color_rgb=(230, 145, 56))
                        else:
                            add_heading(doc, heading_text, level=1, font_size=16, color_rgb=(11, 83, 148))
                    else:
                        if line_text.startswith("-"):
                            add_paragraph(doc, line_text[1:].strip(), style='List Bullet')
                        elif line_text.lstrip().startswith("-"):
                            add_paragraph(doc, line_text.lstrip()[1:].strip(), style='List Bullet 2')
                        else:
                            add_paragraph(doc, line_text)

            # Save the document to a BytesIO buffer
            doc_stream = BytesIO()
            doc.save(doc_stream)
            doc_stream.seek(0)

            try:
                file_metadata = {
                    "name": f"{school_name}.docx",
                    "parents": [school_folder_id]
                }
                media = MediaIoBaseUpload(doc_stream, mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
                uploaded_file = service.files().create(body=file_metadata, media_body=media, fields="id, name, parents").execute()
                uploaded_files.append(uploaded_file)
                print(f"Uploaded {school_name}.docx to Google Drive under {group_name}.")
            except Exception as e:
                print(f"Error uploading file {school_name}.docx: {e}")

    return uploaded_files
